import re
from io import BytesIO

from docx import Document

from ai_translator import (
    translate_with_ai,
)

from final_qa import (
    protect_newlines,
    restore_newlines,
    check_newline_tokens,
    normalize_chinese_punctuation,
    check_final_translation,
    final_qa_summary,
)


# =========================================================
# 韩中游戏本地化 Agent
# 文本 / DOCX / TXT 通用处理层
#
# 核心流程：
#
# 原文
# ↓
# 知识库检索
# ↓
# 唯一正式完整句？
# ├─ 是 → 直接复用，不调用AI
# └─ 否
#      ↓
#   保护真实换行
#      ↓
#   AI受约束翻译
#      ↓
#   删除【需人工确认】
#      ↓
#   检查换行Token
#      ↓
#   恢复真实换行
#      ↓
#   安全中文标点规范化
#      ↓
#   final_qa
#      ↓
#   返回最终业务结果
#
#
# 本模块保持兼容：
#
# xlsx_translator.py
# strict_reviewer.py
# regression_tests.py
# app.py
#
# =========================================================


# =========================================================
# 韩文检测
# =========================================================

KOREAN_PATTERN = re.compile(
    r"[\u1100-\u11FF"
    r"\u3130-\u318F"
    r"\uA960-\uA97F"
    r"\uAC00-\uD7AF"
    r"\uD7B0-\uD7FF]"
)


def contains_korean(
    text
):

    text = str(
        text
        or
        ""
    )

    return bool(
        KOREAN_PATTERN.search(
            text
        )
    )


# =========================================================
# 批量AI致命错误
# =========================================================

class BatchAIUnavailableError(
    Exception
):

    def __init__(
        self,
        message,
        location="",
        error_type="",
    ):

        super().__init__(
            message
        )

        self.location = str(
            location
            or
            ""
        )

        self.error_type = str(
            error_type
            or
            ""
        )


# =========================================================
# 批量处理遇到致命API错误立即停止
# =========================================================

def raise_if_fatal_ai_error(
    result,
    location="",
):

    if not result:
        return


    if not result.get(
        "致命错误",
        False,
    ):

        return


    error_type = str(
        result.get(
            "错误类型",
            ""
        )
        or
        "AIError"
    )


    error_message = str(
        result.get(
            "错误",
            ""
        )
        or
        result.get(
            "错误信息",
            ""
        )
        or
        "AI模型调用发生致命错误。"
    )


    raise BatchAIUnavailableError(
        message=
            error_message,

        location=
            location,

        error_type=
            error_type,
    )


# =========================================================
# 文本工具
# =========================================================

def clean_text(
    value
):

    if value is None:
        return ""

    return str(
        value
    )


# =========================================================
# 保留文本最外层空白
# =========================================================

def split_outer_whitespace(
    text
):

    text = clean_text(
        text
    )


    if not text:

        return (
            "",
            "",
            "",
        )


    match = re.match(
        r"^(\s*)(.*?)(\s*)$",
        text,
        flags=re.DOTALL,
    )


    if not match:

        return (
            "",
            text,
            "",
        )


    return (
        match.group(
            1
        ),
        match.group(
            2
        ),
        match.group(
            3
        ),
    )


# =========================================================
# 去重正式译文
# =========================================================

def unique_translations(
    records
):

    result = []


    for record in records or []:

        chinese = str(
            record.get(
                "msgstr[0]",
                "",
            )
            or
            ""
        ).strip()


        if (
            chinese
            and
            chinese not in result
        ):

            result.append(
                chinese
            )


    return result


# =========================================================
# 检查正式译文是否违反角色名库
# =========================================================

def role_conflicts_in_translation(
    translation,
    search_result,
):

    translation = str(
        translation
        or
        ""
    )


    conflicts = []


    for role in (
        search_result.get(
            "角色名",
            [],
        )
        if search_result
        else []
    ):

        korean_name = str(
            role.get(
                "韩文角色名",
                "",
            )
            or
            ""
        ).strip()


        official_name = str(
            role.get(
                "正式中文名",
                "",
            )
            or
            ""
        ).strip()


        if (
            korean_name
            and
            official_name
            and
            official_name not in translation
        ):

            conflicts.append(
                (
                    f"{korean_name}必须使用"
                    f"正式角色名“{official_name}”"
                )
            )


    return conflicts


# =========================================================
# 正式完整句决定
#
# 优先级：
#
# UWO ＞ Quest
# =========================================================

def choose_exact_translation(
    search_result
):

    search_result = (
        search_result
        or
        {}
    )


    uwo_records = (
        search_result.get(
            "UWO完整句",
            [],
        )
        or
        []
    )


    quest_records = (
        search_result.get(
            "Quest完整句",
            [],
        )
        or
        []
    )


    # =====================================================
    # UWO
    # =====================================================

    if uwo_records:

        translations = (
            unique_translations(
                uwo_records
            )
        )


        if len(
            translations
        ) == 1:

            translation = (
                translations[
                    0
                ]
            )


            role_conflicts = (
                role_conflicts_in_translation(
                    translation,
                    search_result,
                )
            )


            if role_conflicts:

                return {
                    "可直接复用":
                        False,

                    "译文":
                        translation,

                    "来源":
                        "UWO正式主库",

                    "存在风险":
                        True,

                    "风险原因":
                        "；".join(
                            role_conflicts
                        ),

                    "候选译文":
                        translations,
                }


            return {
                "可直接复用":
                    True,

                "译文":
                    translation,

                "来源":
                    "UWO正式主库",

                "存在风险":
                    False,

                "风险原因":
                    "",

                "候选译文":
                    translations,
            }


        if len(
            translations
        ) > 1:

            return {
                "可直接复用":
                    False,

                "译文":
                    "",

                "来源":
                    "UWO正式主库",

                "存在风险":
                    True,

                "风险原因":
                    "UWO正式主库完整句存在多个不同历史译文。",

                "候选译文":
                    translations,
            }


    # =====================================================
    # Quest
    # =====================================================

    if quest_records:

        translations = (
            unique_translations(
                quest_records
            )
        )


        if len(
            translations
        ) == 1:

            translation = (
                translations[
                    0
                ]
            )


            role_conflicts = (
                role_conflicts_in_translation(
                    translation,
                    search_result,
                )
            )


            if role_conflicts:

                return {
                    "可直接复用":
                        False,

                    "译文":
                        translation,

                    "来源":
                        "Quest正式主库",

                    "存在风险":
                        True,

                    "风险原因":
                        "；".join(
                            role_conflicts
                        ),

                    "候选译文":
                        translations,
                }


            return {
                "可直接复用":
                    True,

                "译文":
                    translation,

                "来源":
                    "Quest正式主库",

                "存在风险":
                    False,

                "风险原因":
                    "",

                "候选译文":
                    translations,
            }


        if len(
            translations
        ) > 1:

            return {
                "可直接复用":
                    False,

                "译文":
                    "",

                "来源":
                    "Quest正式主库",

                "存在风险":
                    True,

                "风险原因":
                    "Quest正式主库完整句存在多个不同历史译文。",

                "候选译文":
                    translations,
            }


    return {
        "可直接复用":
            False,

        "译文":
            "",

        "来源":
            "",

        "存在风险":
            False,

        "风险原因":
            "",

        "候选译文":
            [],
    }


# =========================================================
# 构建前后文
# =========================================================

def build_neighbor_context(
    previous_text="",
    next_text="",
    extra_context="",
):

    parts = []


    previous_text = str(
        previous_text
        or
        ""
    ).strip()


    next_text = str(
        next_text
        or
        ""
    ).strip()


    extra_context = str(
        extra_context
        or
        ""
    ).strip()


    if previous_text:

        parts.append(
            (
                "【上一条】\n"
                +
                previous_text
            )
        )


    if next_text:

        parts.append(
            (
                "【下一条】\n"
                +
                next_text
            )
        )


    if extra_context:

        parts.append(
            (
                "【补充上下文】\n"
                +
                extra_context
            )
        )


    return "\n\n".join(
        parts
    )


# =========================================================
# AI人工确认标记
# =========================================================

MANUAL_CONFIRM_PATTERN = re.compile(
    r"【需人工确认\s*[:：]\s*(.*?)】",
    flags=re.DOTALL,
)


def ai_marked_manual_confirmation(
    text
):

    text = str(
        text
        or
        ""
    )


    matches = (
        MANUAL_CONFIRM_PATTERN.findall(
            text
        )
    )


    reasons = []


    for match in matches:

        reason = str(
            match
            or
            ""
        ).strip()


        if (
            reason
            and
            reason not in reasons
        ):

            reasons.append(
                reason
            )


    return (
        len(
            matches
        )
        > 0,

        "；".join(
            reasons
        ),
    )


# =========================================================
# 删除人工确认标记
#
# 标记进入审计字段，
# 不进入正式中文。
# =========================================================

def remove_manual_confirmation_marker(
    text
):

    text = str(
        text
        or
        ""
    )


    has_marker, reason = (
        ai_marked_manual_confirmation(
            text
        )
    )


    cleaned = (
        MANUAL_CONFIRM_PATTERN.sub(
            "",
            text,
        )
    )


    cleaned = (
        cleaned.rstrip()
    )


    return (
        cleaned,
        has_marker,
        reason,
    )


# =========================================================
# 合并人工确认原因
# =========================================================

def merge_reasons(
    *values
):

    result = []


    for value in values:

        if isinstance(
            value,
            list,
        ):

            candidates = value

        else:

            candidates = [
                value
            ]


        for candidate in candidates:

            text = str(
                candidate
                or
                ""
            ).strip()


            if (
                text
                and
                text not in result
            ):

                result.append(
                    text
                )


    return "；".join(
        result
    )


# =========================================================
# 检查意外新增的换行Token
# =========================================================

NEWLINE_TOKEN_PATTERN = re.compile(
    r"⟦NL_(\d+)⟧"
)


def find_unexpected_newline_tokens(
    text,
    expected_count,
):

    text = str(
        text
        or
        ""
    )


    unexpected = []


    for match in (
        NEWLINE_TOKEN_PATTERN.finditer(
            text
        )
    ):

        try:

            index = int(
                match.group(
                    1
                )
            )

        except Exception:

            continue


        if (
            index < 1
            or
            index > expected_count
        ):

            token = (
                match.group(
                    0
                )
            )


            if token not in unexpected:

                unexpected.append(
                    token
                )


    return unexpected


# =========================================================
# 统一文本处理结果
# =========================================================

def make_translation_result(
    source_text,
    translation,
    method,
    source="",
    format_check="",
    need_manual=False,
    confirmation_reason="",
    ai_success=False,
    error_type="",
    error_message="",
    fatal=False,
    error_stage="",
    elapsed=None,
    model="",
    final_qa_passed=None,
    punctuation_normalized=False,
    punctuation_changes=None,
):

    return {
        "原文":
            source_text,

        "译文":
            translation,

        "处理方式":
            method,

        "来源":
            source,

        "格式检查":
            format_check,

        "最终QA通过":
            final_qa_passed,

        "标点自动规范化":
            bool(
                punctuation_normalized
            ),

        "标点规范化内容":
            "；".join(
                punctuation_changes
                or
                []
            ),

        "需人工确认":
            bool(
                need_manual
            ),

        "确认原因":
            str(
                confirmation_reason
                or
                ""
            ),

        "AI成功":
            bool(
                ai_success
            ),

        "错误类型":
            str(
                error_type
                or
                ""
            ),

        "错误阶段":
            str(
                error_stage
                or
                ""
            ),

        "错误":
            str(
                error_message
                or
                ""
            ),

        "错误信息":
            str(
                error_message
                or
                ""
            ),

        "致命错误":
            bool(
                fatal
            ),

        "耗时秒":
            elapsed,

        "模型":
            str(
                model
                or
                ""
            ),
    }


# =========================================================
# 核心：
# 单个文本单元翻译
# =========================================================

def translate_text_unit(
    text,
    knowledge_base,
    extra_context="",
):

    original_text = str(
        text
        or
        ""
    )


    # =====================================================
    # 空文本
    # =====================================================

    if not original_text:

        return make_translation_result(
            source_text=
                original_text,

            translation=
                original_text,

            method=
                "空文本跳过",

            format_check=
                "未执行",

            final_qa_passed=
                True,
        )


    # =====================================================
    # 纯空白
    # =====================================================

    if not original_text.strip():

        return make_translation_result(
            source_text=
                original_text,

            translation=
                original_text,

            method=
                "纯空白跳过",

            format_check=
                "未执行",

            final_qa_passed=
                True,
        )


    # =====================================================
    # 没有韩文
    # =====================================================

    if not contains_korean(
        original_text
    ):

        return make_translation_result(
            source_text=
                original_text,

            translation=
                original_text,

            method=
                "非韩文跳过",

            format_check=
                "未执行",

            final_qa_passed=
                True,
        )


    # =====================================================
    # 保留最外层空白
    # =====================================================

    (
        leading_ws,
        core_text,
        trailing_ws,
    ) = split_outer_whitespace(
        original_text
    )


    if not core_text:

        return make_translation_result(
            source_text=
                original_text,

            translation=
                original_text,

            method=
                "纯空白跳过",

            format_check=
                "未执行",

            final_qa_passed=
                True,
        )


    # =====================================================
    # 查询知识库
    # =====================================================

    try:

        search_result = (
            knowledge_base.search(
                core_text
            )
        )


    except Exception as error:

        return make_translation_result(
            source_text=
                original_text,

            translation=
                "",

            method=
                "知识库查询失败",

            source=
                "",

            format_check=
                "未执行",

            need_manual=
                True,

            confirmation_reason=
                "知识库查询失败。",

            ai_success=
                False,

            error_type=
                type(
                    error
                ).__name__,

            error_message=
                str(
                    error
                ),

            fatal=
                True,

            error_stage=
                "知识库查询阶段",

            final_qa_passed=
                False,
        )


    # =====================================================
    # 唯一正式完整句
    #
    # 不调用AI。
    # =====================================================

    exact_decision = (
        choose_exact_translation(
            search_result
        )
    )


    if exact_decision.get(
        "可直接复用",
        False,
    ):

        translation_core = str(
            exact_decision.get(
                "译文",
                ""
            )
            or
            ""
        )


        # =================================================
        # 正式完整句也执行安全标点规范化
        #
        # 但不会改变术语、语义、数字等。
        # =================================================

        punctuation_result = (
            normalize_chinese_punctuation(
                translation_core
            )
        )


        normalized_translation = (
            punctuation_result.get(
                "文本",
                translation_core,
            )
        )


        qa_result = (
            check_final_translation(
                core_text,
                normalized_translation,
            )
        )


        qa_summary = (
            final_qa_summary(
                core_text,
                normalized_translation,
            )
        )


        need_manual = not (
            qa_result.get(
                "通过",
                False,
            )
        )


        confirmation_reason = ""


        if need_manual:

            confirmation_reason = (
                (
                    "唯一正式完整句已直接复用，"
                    "但最终QA发现风险："
                )
                +
                qa_summary
            )


        final_translation = (
            leading_ws
            +
            normalized_translation
            +
            trailing_ws
        )


        return make_translation_result(
            source_text=
                original_text,

            translation=
                final_translation,

            method=
                "正式完整句直接复用",

            source=
                exact_decision.get(
                    "来源",
                    "",
                ),

            format_check=
                qa_summary,

            need_manual=
                need_manual,

            confirmation_reason=
                confirmation_reason,

            ai_success=
                False,

            fatal=
                False,

            final_qa_passed=
                qa_result.get(
                    "通过",
                    False,
                ),

            punctuation_normalized=
                punctuation_result.get(
                    "是否修改",
                    False,
                ),

            punctuation_changes=
                punctuation_result.get(
                    "修改项",
                    [],
                ),
        )


    # =====================================================
    # AI翻译前：
    # 保护真实换行
    # =====================================================

    (
        protected_source,
        newline_types,
    ) = protect_newlines(
        core_text
    )


    combined_context = str(
        extra_context
        or
        ""
    ).strip()


    # =====================================================
    # 有真实换行时，
    # 告诉AI换行Token禁止修改。
    # =====================================================

    if newline_types:

        newline_instruction = (
            "特别格式要求："
            "韩文原文中的⟦NL_1⟧、⟦NL_2⟧等"
            "是系统用于保护原始真实换行的位置标记。"
            "必须原样、原数量、原顺序保留，"
            "不得删除、移动、翻译、合并或新增。"
        )


        if combined_context:

            combined_context = (
                combined_context
                +
                "\n\n"
                +
                newline_instruction
            )

        else:

            combined_context = (
                newline_instruction
            )


    # =====================================================
    # 知识库冲突
    # =====================================================

    knowledge_risk = bool(
        exact_decision.get(
            "存在风险",
            False,
        )
    )


    knowledge_risk_reason = str(
        exact_decision.get(
            "风险原因",
            ""
        )
        or
        ""
    ).strip()


    # =====================================================
    # 调用AI
    # =====================================================

    ai_result = (
        translate_with_ai(
            korean_text=
                protected_source,

            search_result=
                search_result,

            extra_context=
                combined_context,

            mode=
                "快速翻译",
        )
    )


    # =====================================================
    # AI失败
    # =====================================================

    if not ai_result.get(
        "成功",
        False,
    ):

        error_message = str(
            ai_result.get(
                "错误",
                ""
            )
            or
            ai_result.get(
                "错误信息",
                ""
            )
            or
            "AI模型未能完成翻译。"
        )


        return make_translation_result(
            source_text=
                original_text,

            translation=
                original_text,

            method=
                "AI调用失败，保留韩文",

            source=
                "AI模型",

            format_check=
                "未执行",

            need_manual=
                True,

            confirmation_reason=
                error_message,

            ai_success=
                False,

            error_type=
                ai_result.get(
                    "错误类型",
                    "",
                ),

            error_message=
                error_message,

            fatal=
                ai_result.get(
                    "致命错误",
                    False,
                ),

            error_stage=
                ai_result.get(
                    "错误阶段",
                    "",
                ),

            elapsed=
                ai_result.get(
                    "耗时秒"
                ),

            model=
                ai_result.get(
                    "模型",
                    "",
                ),

            final_qa_passed=
                False,
        )


    # =====================================================
    # AI成功
    # =====================================================

    raw_translation = str(
        ai_result.get(
            "译文",
            ""
        )
        or
        ""
    )


    # =====================================================
    # 第一步：
    # 删除【需人工确认】
    #
    # QA不能针对带标记的AI原始回答。
    # =====================================================

    (
        cleaned_translation,
        ai_manual,
        ai_manual_reason,
    ) = remove_manual_confirmation_marker(
        raw_translation
    )


    # =====================================================
    # 第二步：
    # 检查换行Token是否被删除
    # =====================================================

    token_result = (
        check_newline_tokens(
            cleaned_translation,
            newline_types,
        )
    )


    missing_tokens = (
        token_result.get(
            "缺失Token",
            [],
        )
    )


    unexpected_tokens = (
        find_unexpected_newline_tokens(
            cleaned_translation,
            len(
                newline_types
            ),
        )
    )


    # =====================================================
    # 第三步：
    # 恢复真实换行
    # =====================================================

    restored_translation = (
        restore_newlines(
            cleaned_translation,
            newline_types,
        )
    )


    # =====================================================
    # 第四步：
    # 安全中文标点自动规范化
    #
    # 例如：
    #
    # …   → ……
    # ... → ……
    # ?   → ？
    # !   → ！
    # ?!  → ？！
    # !?  → ！？
    # ,   → ，
    # ;   → ；
    # :   → ：
    #
    # URL / Email / 标签 / 变量 / 数字格式等
    # 由final_qa.py保护。
    # =====================================================

    punctuation_result = (
        normalize_chinese_punctuation(
            restored_translation
        )
    )


    normalized_translation = (
        punctuation_result.get(
            "文本",
            restored_translation,
        )
    )


    # =====================================================
    # 第五步：
    # 对真正准备写入文件的最终中文执行final_qa
    # =====================================================

    qa_result = (
        check_final_translation(
            core_text,
            normalized_translation,
        )
    )


    qa_summary = (
        final_qa_summary(
            core_text,
            normalized_translation,
        )
    )


    # =====================================================
    # 人工确认原因
    # =====================================================

    manual_reasons = []


    if knowledge_risk:

        manual_reasons.append(
            knowledge_risk_reason
            or
            "正式知识库存在多个历史译法或冲突。"
        )


    if ai_manual:

        manual_reasons.append(
            ai_manual_reason
            or
            "AI判断需要人工确认。"
        )


    if missing_tokens:

        manual_reasons.append(
            (
                "AI删除了受保护的真实换行标记："
                +
                ", ".join(
                    missing_tokens
                )
            )
        )


    if unexpected_tokens:

        manual_reasons.append(
            (
                "AI新增了不存在的换行标记："
                +
                ", ".join(
                    unexpected_tokens
                )
            )
        )


    if not qa_result.get(
        "通过",
        False,
    ):

        manual_reasons.append(
            qa_summary
        )


    need_manual = bool(
        manual_reasons
    )


    confirmation_reason = (
        merge_reasons(
            manual_reasons
        )
    )


    # =====================================================
    # 最终准备写入文件的中文
    # =====================================================

    final_translation = (
        leading_ws
        +
        normalized_translation
        +
        trailing_ws
    )


    return make_translation_result(
        source_text=
            original_text,

        translation=
            final_translation,

        method=
            (
                "AI受约束翻译-需人工确认"
                if need_manual
                else
                "AI受约束翻译"
            ),

        source=
            "AI模型",

        format_check=
            qa_summary,

        need_manual=
            need_manual,

        confirmation_reason=
            confirmation_reason,

        ai_success=
            True,

        error_type=
            "",

        error_message=
            "",

        fatal=
            False,

        error_stage=
            "",

        elapsed=
            ai_result.get(
                "耗时秒"
            ),

        model=
            ai_result.get(
                "模型",
                "",
            ),

        final_qa_passed=
            (
                qa_result.get(
                    "通过",
                    False,
                )
                and
                not missing_tokens
                and
                not unexpected_tokens
            ),

        punctuation_normalized=
            punctuation_result.get(
                "是否修改",
                False,
            ),

        punctuation_changes=
            punctuation_result.get(
                "修改项",
                [],
            ),
    )


# =========================================================
# DOCX/TXT统计
# =========================================================

def new_document_stats():

    return {
        "文本单元总数":
            0,

        "韩文文本数":
            0,

        "正式完整句直接复用":
            0,

        "AI翻译成功":
            0,

        "AI翻译失败":
            0,

        "需人工确认":
            0,

        "最终QA异常":
            0,

        "标点自动规范化":
            0,

        "非韩文跳过":
            0,

        "空文本跳过":
            0,
    }


# =========================================================
# 更新统计
# =========================================================

def update_document_stats(
    stats,
    result,
):

    stats[
        "文本单元总数"
    ] += 1


    method = str(
        result.get(
            "处理方式",
            ""
        )
        or
        ""
    )


    if method in {
        "空文本跳过",
        "纯空白跳过",
    }:

        stats[
            "空文本跳过"
        ] += 1

        return


    if method == "非韩文跳过":

        stats[
            "非韩文跳过"
        ] += 1

        return


    stats[
        "韩文文本数"
    ] += 1


    if method == "正式完整句直接复用":

        stats[
            "正式完整句直接复用"
        ] += 1


    if result.get(
        "AI成功",
        False,
    ):

        stats[
            "AI翻译成功"
        ] += 1


    if (
        method
        ==
        "AI调用失败，保留韩文"
    ):

        stats[
            "AI翻译失败"
        ] += 1


    if result.get(
        "需人工确认",
        False,
    ):

        stats[
            "需人工确认"
        ] += 1


    if result.get(
        "最终QA通过"
    ) is False:

        stats[
            "最终QA异常"
        ] += 1


    if result.get(
        "标点自动规范化",
        False,
    ):

        stats[
            "标点自动规范化"
        ] += 1


# =========================================================
# 明细记录
# =========================================================

def make_detail_record(
    location,
    result
):

    return {
        "位置":
            location,

        "原文":
            result.get(
                "原文",
                "",
            ),

        "译文":
            result.get(
                "译文",
                "",
            ),

        "处理方式":
            result.get(
                "处理方式",
                "",
            ),

        "来源":
            result.get(
                "来源",
                "",
            ),

        "格式检查":
            result.get(
                "格式检查",
                "",
            ),

        "最终QA通过":
            result.get(
                "最终QA通过"
            ),

        "标点自动规范化":
            result.get(
                "标点自动规范化",
                False,
            ),

        "标点规范化内容":
            result.get(
                "标点规范化内容",
                "",
            ),

        "需人工确认":
            result.get(
                "需人工确认",
                False,
            ),

        "确认原因":
            result.get(
                "确认原因",
                "",
            ),

        "AI成功":
            result.get(
                "AI成功",
                False,
            ),

        "错误类型":
            result.get(
                "错误类型",
                "",
            ),

        "错误阶段":
            result.get(
                "错误阶段",
                "",
            ),

        "错误":
            result.get(
                "错误",
                "",
            ),

        "耗时秒":
            result.get(
                "耗时秒"
            ),

        "模型":
            result.get(
                "模型",
                "",
            ),
    }


# =========================================================
# DOCX表格
# =========================================================

def process_table(
    table,
    knowledge_base,
    stats,
    details,
    table_index=1,
):

    for row_index, row in enumerate(
        table.rows,
        start=1,
    ):

        for cell_index, cell in enumerate(
            row.cells,
            start=1,
        ):

            paragraphs = list(
                cell.paragraphs
            )


            for paragraph_index, paragraph in enumerate(
                paragraphs,
                start=1,
            ):

                text = (
                    paragraph.text
                )


                if not text:
                    continue


                previous_text = ""


                if paragraph_index > 1:

                    previous_text = (
                        paragraphs[
                            paragraph_index
                            -
                            2
                        ].text
                    )


                next_text = ""


                if paragraph_index < len(
                    paragraphs
                ):

                    next_text = (
                        paragraphs[
                            paragraph_index
                        ].text
                    )


                context = (
                    build_neighbor_context(
                        previous_text=
                            previous_text,

                        next_text=
                            next_text,
                    )
                )


                result = (
                    translate_text_unit(
                        text=
                            text,

                        knowledge_base=
                            knowledge_base,

                        extra_context=
                            context,
                    )
                )


                location = (
                    f"表格{table_index}"
                    f"-第{row_index}行"
                    f"-第{cell_index}列"
                    f"-段落{paragraph_index}"
                )


                raise_if_fatal_ai_error(
                    result=
                        result,

                    location=
                        location,
                )


                update_document_stats(
                    stats,
                    result,
                )


                details.append(
                    make_detail_record(
                        location,
                        result,
                    )
                )


                if (
                    result.get(
                        "处理方式"
                    )
                    ==
                    "AI调用失败，保留韩文"
                ):

                    continue


                translation = (
                    result.get(
                        "译文",
                        text,
                    )
                )


                if translation != text:

                    paragraph.text = (
                        translation
                    )


            # =============================================
            # 嵌套表格
            # =============================================

            for nested_index, nested_table in enumerate(
                cell.tables,
                start=1,
            ):

                process_table(
                    table=
                        nested_table,

                    knowledge_base=
                        knowledge_base,

                    stats=
                        stats,

                    details=
                        details,

                    table_index=
                        (
                            f"{table_index}"
                            f".{row_index}"
                            f".{cell_index}"
                            f".{nested_index}"
                        ),
                )


# =========================================================
# DOCX
# =========================================================

def process_docx(
    file_bytes,
    knowledge_base,
):

    document = Document(
        BytesIO(
            file_bytes
        )
    )


    stats = (
        new_document_stats()
    )


    details = []


    paragraphs = list(
        document.paragraphs
    )


    # =====================================================
    # 普通段落
    # =====================================================

    for index, paragraph in enumerate(
        paragraphs
    ):

        text = (
            paragraph.text
        )


        if not text:
            continue


        previous_text = (
            paragraphs[
                index - 1
            ].text
            if index > 0
            else
            ""
        )


        next_text = (
            paragraphs[
                index + 1
            ].text
            if index + 1
            <
            len(
                paragraphs
            )
            else
            ""
        )


        context = (
            build_neighbor_context(
                previous_text=
                    previous_text,

                next_text=
                    next_text,
            )
        )


        result = (
            translate_text_unit(
                text=
                    text,

                knowledge_base=
                    knowledge_base,

                extra_context=
                    context,
            )
        )


        location = (
            f"DOCX正文第{index + 1}段"
        )


        raise_if_fatal_ai_error(
            result=
                result,

            location=
                location,
        )


        update_document_stats(
            stats,
            result,
        )


        details.append(
            make_detail_record(
                location,
                result,
            )
        )


        if (
            result.get(
                "处理方式"
            )
            ==
            "AI调用失败，保留韩文"
        ):

            continue


        translation = (
            result.get(
                "译文",
                text,
            )
        )


        if translation != text:

            paragraph.text = (
                translation
            )


    # =====================================================
    # 表格
    # =====================================================

    for table_index, table in enumerate(
        document.tables,
        start=1,
    ):

        process_table(
            table=
                table,

            knowledge_base=
                knowledge_base,

            stats=
                stats,

            details=
                details,

            table_index=
                table_index,
        )


    output = BytesIO()


    document.save(
        output
    )


    output.seek(
        0
    )


    return (
        output.getvalue(),
        stats,
        details,
    )


# =========================================================
# TXT
# =========================================================

def process_txt(
    file_bytes,
    knowledge_base,
):

    text = None


    # =====================================================
    # 编码识别
    # =====================================================

    for encoding in [
        "utf-8-sig",
        "utf-8",
        "cp949",
        "euc-kr",
    ]:

        try:

            text = file_bytes.decode(
                encoding
            )

            break


        except Exception:

            continue


    if text is None:

        raise ValueError(
            "无法识别TXT文件编码。"
        )


    lines = text.splitlines(
        keepends=True
    )


    stats = (
        new_document_stats()
    )


    details = []

    output_lines = []


    for index, line in enumerate(
        lines
    ):

        # =================================================
        # 分离TXT物理行尾
        # =================================================

        line_ending = ""


        if line.endswith(
            "\r\n"
        ):

            body = line[
                :-2
            ]

            line_ending = "\r\n"


        elif line.endswith(
            "\n"
        ):

            body = line[
                :-1
            ]

            line_ending = "\n"


        elif line.endswith(
            "\r"
        ):

            body = line[
                :-1
            ]

            line_ending = "\r"


        else:

            body = line


        previous_text = ""


        if index > 0:

            previous_text = (
                lines[
                    index - 1
                ].rstrip(
                    "\r\n"
                )
            )


        next_text = ""


        if (
            index + 1
            <
            len(
                lines
            )
        ):

            next_text = (
                lines[
                    index + 1
                ].rstrip(
                    "\r\n"
                )
            )


        context = (
            build_neighbor_context(
                previous_text=
                    previous_text,

                next_text=
                    next_text,
            )
        )


        result = (
            translate_text_unit(
                text=
                    body,

                knowledge_base=
                    knowledge_base,

                extra_context=
                    context,
            )
        )


        location = (
            f"TXT第{index + 1}行"
        )


        raise_if_fatal_ai_error(
            result=
                result,

            location=
                location,
        )


        update_document_stats(
            stats,
            result,
        )


        details.append(
            make_detail_record(
                location,
                result,
            )
        )


        if (
            result.get(
                "处理方式"
            )
            ==
            "AI调用失败，保留韩文"
        ):

            translated_body = (
                body
            )

        else:

            translated_body = (
                result.get(
                    "译文",
                    body,
                )
            )


        output_lines.append(
            translated_body
            +
            line_ending
        )


    output_text = "".join(
        output_lines
    )


    return (
        output_text.encode(
            "utf-8-sig"
        ),
        stats,
        details,
    )


# =========================================================
# 文档统一入口
# =========================================================

def process_document(
    file_name,
    file_bytes,
    knowledge_base,
):

    file_name = str(
        file_name
        or
        ""
    ).lower()


    if file_name.endswith(
        ".docx"
    ):

        return process_docx(
            file_bytes=
                file_bytes,

            knowledge_base=
                knowledge_base,
        )


    if file_name.endswith(
        ".txt"
    ):

        return process_txt(
            file_bytes=
                file_bytes,

            knowledge_base=
                knowledge_base,
        )


    raise ValueError(
        (
            "当前document_processor只支持"
            "DOCX和TXT。"
            "XLSX由xlsx_translator.py处理。"
        )
    )


# =========================================================
# 独立加载测试
# =========================================================

if __name__ == "__main__":

    print(
        "DOCX/TXT翻译处理层加载成功。"
    )

    print(
        "已启用："
    )

    print(
        "1. 真实换行强制保护"
    )

    print(
        "2. 人工确认标记先清理再QA"
    )

    print(
        "3. 安全中文标点自动规范化"
    )

    print(
        "4. 韩文残留检测"
    )

    print(
        "5. 最终译文QA"
    )

    print(
        "6. 批量致命API错误立即停止"
    )